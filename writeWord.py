# coding:utf-8
# 写word文档文件
import sys,math
import importlib
from win32com.client import Dispatch
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
def makeImgExcell(doc, allxls,label,cols):
    num = math.ceil(len(allxls) / cols)
    tables = doc.add_table(rows=2 * num, cols=cols)
    for x in range(num * 2):
        if x % 2 == 0:
            for i in range(cols):

                run = tables.cell(x, i).paragraphs[0].add_run()
                try:
                    print("add_picture",int(i + (x / 2)), cols)
                    print("add_picture", allxls[int(i + (x / 2) * cols)])
                    run.add_picture(allxls[int(i + (x / 2) * cols)], width=Inches(2), height=Inches(2.5))

                except IndexError:
                    continue
                    # jpg_ima = Image.open(allxls[int(i + (x / 2) * cols)])
                    # jpg_ima.save(label[int(i + (x / 2) * cols)].split("\\")[-1])
                    # run.add_picture(allxls[int(i + (x / 2) * cols)].split("\\")[-1], width=Inches(2),
                    #                 height=Inches(2.5))
                except Exception:
                    continue


        else:
            for i in range(cols):
                try:
                    tables.cell(x, i).text = label[i + int(x / 2) * cols]
                except:
                    continue
    return doc
def WriteWord(Subject,RemarkImgPaths,resultPaths,confs = []):
    print("start write word report:",Subject,RemarkImgPaths,resultPaths,confs)
    importlib.reload(sys)
    # 创建文档对象
    document = Document()
    #添加页眉
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "ScienceAI"
    # 设置文档标题，中文要用unicode字符串
    fileName = Subject.split("/")[-1]
    document.add_heading(u'ScienceAI %s: Recognition Report ' % fileName, 0)
    # 往文档中添加段落
    p = document.add_paragraph('I identified the document of  ')
    p.add_run('"%s.pdf"'%fileName).bold = True
    p.add_run(',the ')
    p.add_run('result ').italic = True
    p.add_run('as flows:')
    # 添加一级标题
    document.add_heading(u'The PDF document consists of %s pages, each of which is shown below:'% len(RemarkImgPaths), level=1)
    originlabels = []
    originPaths = []
    for i, RemarkImgPath in enumerate(RemarkImgPaths):
        originlabels.append('The page  %s of the PDF' % (i + 1))
        originPaths.append(RemarkImgPath.replace("/result","").replace("-Remarked","") )
    document = makeImgExcell(document, originPaths, originlabels, 3)

    # 添加二级标题
    document.add_heading(u'The graph found is as follows: ', level=2)
    for i, RemarkImgPath in enumerate(RemarkImgPaths):
        document.add_heading('The page  %s of the PDF,found graph ' % (i + 1), level=3)
        # 添加图片，并指定宽度
        document.add_picture(RemarkImgPath, width=Inches(5))

        for j in range(len(resultPaths[i])):
            path = resultPaths[i][j]
            print(path)
            # 添加有序列表
            document.add_paragraph('The page  {} of the PDF,{} graph as follow:' .format((i + 1),(j+1)), style='ListNumber')
            document.add_paragraph('See the detailed data file "%s"!' % path.split("/")[-1].replace("-Seeked.jpg", "-RedrawA(B).txt"), style='IntenseQuote')


            document = makeImgExcell(document, [path,path.replace("-Seeked", "-RedrawPic")],["Seeked Graph","Redraw Graph"], 2)


    #

    # document.add_paragraph('second item in ordered list', style='ListNumber')
    # document.add_paragraph('third item in ordered list', style='ListNumber')

    # 添加表格: 1行3列
    table = document.add_table(rows=1, cols=3)
    #加入边框
    table.style = 'ColorfulShading'
    # 获取第一行的单元格列表对象
    hdr_cells = table.rows[0].cells
    # 为每一个单元格赋值
    # 注：值都要为字符串类型
    hdr_cells[0].text = 'Page'
    hdr_cells[1].text = 'Number of graph'
    hdr_cells[2].text = 'Confidence'
    for i in range(len(RemarkImgPaths)):
        # 为表格添加一行
        new_cells = table.add_row().cells
        new_cells[0].text = str(i+1)
        new_cells[1].text = str(len(resultPaths[i]))
        new_cells[2].text = '%s'%confs[i]
    # 添加分页符
    document.add_page_break()
    # 往新的一页中添加段落
    p = document.add_paragraph('Thank you for your trust. I look forward to seeing you next time.')

    docPath = Subject + '/Report.docx'
    # 保存文档
    document.save(docPath)
    # 保存PDF文档

    # doc2pdf(Subject + '-Report.docx')
    return docPath
def doc2pdf(input_file):
    word = Dispatch('Word.Application')
    doc = word.Documents.Open(input_file)
    doc.SaveAs(input_file.replace(".docx", ".pdf"), FileFormat=17)
    doc.Close()
    word.Quit()

def WriteWordParticle(wordFileName = "demo.docx",originImg = 't0.jpg',markImg1 = 't0.jpg',
                      markImg2 = 't0.jpg', histImg1 = 't1.jpg',histImg2 = 't1.jpg'
                      ):
    # 打开文档
    document = Document()
    # 添加页眉
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "ScienceAI"
    # 加入不同等级的标题
    document.add_heading('Particle image identification report', 0)
    document.add_heading(u'The request file is as follows', 1)
    #document.add_heading(u'二级标题', 2)
    # 增加图片（此处使用相对位置）
    document.add_picture(originImg, width=Inches(1.25))
    document.add_heading(u'We have recognized the picture you sent twice. ',1)
    # 增加无序列表
    document.add_paragraph(u'The dark target recognition result is as follows:',
                           style='List Bullet')
    document.add_paragraph(u'The label images:(t0.jpg)', style='List Number')
    # 增加图片（此处使用相对位置）
    document.add_picture(markImg1, width=Inches(1.25))
    # 增加有序列表
    document.add_paragraph(u'A histogram images:({})'.format(histImg1), style='List Number')
    # 增加图片（此处使用相对位置）
    document.add_picture(histImg1, width=Inches(1.25))
    ######################################第二幅图
    # 增加无序列表
    document.add_paragraph(u'The dark target recognition result is as follows:',
                           style='List Bullet')
    document.add_paragraph(u'The label images:({})'.format(markImg2), style='List Number' )
    # 增加图片（此处使用相对位置）
    document.add_picture(markImg2, width=Inches(1.25))
    # 增加有序列表
    document.add_paragraph(u'A histogram images:({})'.format(histImg2), style='List Number')
    # 增加图片（此处使用相对位置）
    document.add_picture(histImg2, width=Inches(1.25))
    document.add_paragraph(u'All of the above files can be viewed in the zip package ',
                           style='List Bullet')
    document.add_paragraph(u'Thank you for your trust.I look forward to seeing you next time.')
    # 保存文件
    document.save(wordFileName)
if __name__ == '__main__':
    WriteWordParticle()